# from flask import Flask, render_template, request, jsonify
# import os
# import psycopg2
# from transformers import BertTokenizer, BertForQuestionAnswering
# import torch
# import fitz  # PyMuPDF
# from PIL import Image
# import numpy as np
# from docx import Document
# from easyocr import Reader

# app = Flask(__name__)

# # Database connection
# conn = psycopg2.connect(
#     dbname="Document_Upload",
#     user="postgres",
#     password="root",
#     host="localhost",
#     port="5432"
# )

# # # OCR reader

# # Load pre-trained BERT model and tokenizer
# tokenizer = BertTokenizer.from_pretrained("bert-large-uncased-whole-word-masking-finetuned-squad")
# model = BertForQuestionAnswering.from_pretrained("bert-large-uncased-whole-word-masking-finetuned-squad")
# # # OCR reader
# reader = Reader(['en'])
# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload():
#     if 'document' in request.files:
#         documents = request.files.getlist('document')  # Get list of uploaded files
#         for document in documents:
#             if document.filename == '':
#                 continue
#             filename = document.filename
#             filepath = os.path.join('uploads', filename)
#             document.save(filepath)
            
#             # Extract text from the document
#             if filename.lower().endswith('.pdf'):
#                 text = extract_text_from_pdf(filepath)
#             elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
#                 text = extract_text_from_image(filepath)
#             elif filename.lower().endswith('.docx'):  # Add support for .docx files
#                 text = extract_text_from_docx(filepath)
#             else:
#                 return "Unsupported file format"
            
#             # Store text in a text file
#             text_filepath = os.path.splitext(filepath)[0] + ".txt"
#             with open(text_filepath, 'w', encoding='utf-8') as f:
#                 f.write(text)
            
#             # Store file in the database
#             store_file_in_database(filename, text)
        
#         # Combine text from all individual files into one text file
#         combined_text = combine_text_files('uploads')
#         combined_text_filepath = os.path.join('combined_files', 'combined_text.txt')
#         with open(combined_text_filepath, 'w', encoding='utf-8') as f:
#             f.write(combined_text)
        
#         return "Files successfully uploaded and converted to text!"
#     return "No files selected."

# @app.route('/fetch/<int:doc_id>', methods=['GET'])
# def fetch_document(doc_id):
#     cursor = conn.cursor()
#     query = "SELECT filename, content FROM uploaded_files WHERE id = %s"
#     cursor.execute(query, (doc_id,))
#     document = cursor.fetchone()
#     cursor.close()
#     if document:
#         filename, content = document
#         return jsonify({"filename": filename, "content": content})
#     else:
#         return jsonify({"error": "Document not found"}), 404

# # @app.route('/question-answering', methods=['POST'])
# # def question_answering():
# #     data = request.json
# #     question = data.get('question')
# #     doc_id = data.get('doc_id')
    
# #     # Fetch content from the database
# #     cursor = conn.cursor()
# #     query = "SELECT content FROM uploaded_files WHERE id = %s"
# #     cursor.execute(query, (doc_id,))
# #     document_content = cursor.fetchone()[0]
# #     cursor.close()
    
# #     # Use the pre-trained model to generate an answer
# #     inputs = tokenizer(question, document_content, return_tensors="pt", padding=True, truncation=True)
# #     start_positions, end_positions = model(**inputs)
# #     answer_start = torch.argmax(start_positions)
# #     answer_end = torch.argmax(end_positions) + 1
# #     answer = tokenizer.convert_tokens_to_string(tokenizer.convert_ids_to_tokens(inputs["input_ids"][0][answer_start:answer_end]))
    
# #     return jsonify({"answer": answer})
# @app.route('/question-answering', methods=['POST'])
# def question_answering():
#     data = request.json
#     question = data.get('question')
#     doc_id = data.get('doc_id')
    
#     # Fetch content from the database
#     cursor = conn.cursor()
#     query = "SELECT content FROM uploaded_files WHERE id = %s"
#     cursor.execute(query, (doc_id,))
#     document_content = cursor.fetchone()[0]
#     cursor.close()
    
#     # Use the pre-trained model to generate an answer
#     inputs = tokenizer(question, document_content, return_tensors="pt", padding=True, truncation=True)
#     outputs = model(**inputs)
    
#     # Extract start and end positions from the model output
#     start_positions = outputs.start_logits.squeeze().argmax().item()
#     end_positions = outputs.end_logits.squeeze().argmax().item() + 1
    
#     answer = tokenizer.convert_tokens_to_string(tokenizer.convert_ids_to_tokens(inputs["input_ids"][0][start_positions:end_positions]))
    
#     return jsonify({"answer": answer})


# # def extract_text_from_pdf(filepath):
# #     text = ""
# #     doc = fitz.open(filepath)
# #     for page in doc:
# #         text += page.get_text()
# #     return text

# # def extract_text_from_image(filepath):
# #     image = Image.open(filepath)
# #     # Convert image to numpy array
# #     image_np = np.array(image)
# #     # Extract text from image (using OCR)
# #     # Add your OCR code here
# #     return ""

# def extract_text_from_docx(filepath):
#     doc = Document(filepath)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + '\n'
#     return text
# def extract_text_from_pdf(filepath):
#     text = ""
#     doc = fitz.open(filepath)
#     for page in doc:
#         text += page.get_text()
#     return text

# def extract_text_from_image(filepath):
#     image = Image.open(filepath)
#     # Convert image to numpy array
#     image_np = np.array(image)
#     # Extract text from image
#     extracted_text = reader.readtext(image_np)

#     # Check if extracted_text is empty
#     if not extracted_text:
#         return ""

#     # Iterate over each element in extracted_text and extract the text content
#     text_list = []
#     for text_info in extracted_text:
#         text_list.append(text_info[1])

#     # Join the extracted text content into a single string
#     text = '\n'.join(text_list)
#     return text

# # def extract_text_from_docx(filepath):
# #     doc = Document(filepath)
# #     text = ""
# #     for paragraph in doc.paragraphs:
# #         text += paragraph.text + '\n'
# #     return text
# def store_file_in_database(filename, content):
#     cursor = conn.cursor()
#     query = "INSERT INTO uploaded_files (filename, content) VALUES (%s, %s)"
#     cursor.execute(query, (filename, content))
#     conn.commit()
#     cursor.close()

# def combine_text_files(folder):
#     combined_text = ""
#     for root, dirs, files in os.walk(folder):
#         for file in files:
#             if file.endswith('.txt'):
#                 file_path = os.path.join(root, file)
#                 with open(file_path, 'r', encoding='utf-8') as f:
#                     combined_text += f.read() + "\n\n"
#     return combined_text

# if __name__ == '__main__':
#     app.run(debug=True)











from flask import Flask, render_template, request, jsonify
import os
import psycopg2
import fitz  # PyMuPDF
from PIL import Image
import numpy as np
from docx import Document
from easyocr import Reader
import torch
from transformers import BertTokenizer, BertModel

app = Flask(__name__)

# Database connection
conn = psycopg2.connect(
    dbname="Document_Upload",
    user="postgres",
    password="root",
    host="localhost",
    port="5432"
)

# Load pre-trained BERT model and tokenizer
tokenizer = BertTokenizer.from_pretrained("bert-base-uncased")
model = BertModel.from_pretrained("bert-base-uncased")

# OCR reader
reader = Reader(['en'])

def text_to_bert_embeddings(text):
    inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True)
    with torch.no_grad():
        outputs = model(**inputs)
    embeddings = outputs.last_hidden_state.mean(dim=1).squeeze().tolist()
    print("BERT embeddings:", embeddings)  # Print embeddings
    return embeddings

@app.route('/')
def index():
    return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload():
#     if 'document' in request.files:
#         documents = request.files.getlist('document')  # Get list of uploaded files
#         for document in documents:
#             if document.filename == '':
#                 continue
#             filename = document.filename
#             filepath = os.path.join('uploads', filename)
#             document.save(filepath)
            
#             # Extract text from the document
#             if filename.lower().endswith('.pdf'):
#                 text = extract_text_from_pdf(filepath)
#             elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
#                 text = extract_text_from_image(filepath)
#             elif filename.lower().endswith('.docx'):  # Add support for .docx files
#                 text = extract_text_from_docx(filepath)
#             else:
#                 return "Unsupported file format"
            
#             # Convert text to BERT embeddings and display in console
#             embeddings = text_to_bert_embeddings(text)
#             print("BERT embeddings for", filename, ":", embeddings)
            
#             # Store file in the database
#             store_file_in_database(filename, text)
        
#         # Combine text from all individual files into one text file
#         combined_text = combine_text_files('uploads')
#         combined_text_filepath = os.path.join('combined_files', 'combined_text.txt')
#         with open(combined_text_filepath, 'w', encoding='utf-8') as f:
#             f.write(combined_text)
        
#         return "Files successfully uploaded and converted to text!"
#     return "No files selected."

@app.route('/upload', methods=['POST'])
def upload():
    if 'document' in request.files:
        documents = request.files.getlist('document')  # Get list of uploaded files
        for document in documents:
            if document.filename == '':
                continue
            filename = document.filename
            filepath = os.path.join('uploads', filename)
            document.save(filepath)
            
            # Extract text from the document
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
                text = extract_text_from_image(filepath)
            elif filename.lower().endswith('.docx'):  # Add support for .docx files
                text = extract_text_from_docx(filepath)
            else:
                return "Unsupported file format"
            
            # Store text in a text file
            text_filepath = os.path.splitext(filepath)[0] + ".txt"
            with open(text_filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # Store file in the database
            store_file_in_database(filename, text)
        
        # Combine text from all individual files into one text file
        combined_text = combine_text_files('uploads')
        combined_text_filepath = os.path.join('combined_files', 'combined_text.txt')
        with open(combined_text_filepath, 'w', encoding='utf-8') as f:
            f.write(combined_text)
        
        return "Files successfully uploaded and converted to text!"
    return "No files selected."


@app.route('/fetch/<int:doc_id>', methods=['GET'])
def fetch_document(doc_id):
    cursor = conn.cursor()
    query = "SELECT filename, content FROM uploaded_files WHERE id = %s"
    cursor.execute(query, (doc_id,))
    document = cursor.fetchone()
    cursor.close()
    if document:
        filename, content = document
        return jsonify({"filename": filename, "content": content})
    else:
        return jsonify({"error": "Document not found"}), 404

@app.route('/question-answering', methods=['POST'])
def question_answering():
    data = request.json
    question = data.get('question')
    doc_id = data.get('doc_id')
    
    # Fetch content from the database
    cursor = conn.cursor()
    query = "SELECT content FROM uploaded_files WHERE id = %s"
    cursor.execute(query, (doc_id,))
    document_content = cursor.fetchone()[0]
    cursor.close()
    
    # Use the pre-trained model to generate an answer
    inputs = tokenizer(question, document_content, return_tensors="pt", padding=True, truncation=True)
    outputs = model(**inputs)
    
    # Extract start and end positions from the model output
    start_positions = outputs.start_logits.squeeze().argmax().item()
    end_positions = outputs.end_logits.squeeze().argmax().item() + 1
    
    answer = tokenizer.convert_tokens_to_string(tokenizer.convert_ids_to_tokens(inputs["input_ids"][0][start_positions:end_positions]))
    
    return jsonify({"answer": answer})

def extract_text_from_pdf(filepath):
    text = ""
    doc = fitz.open(filepath)
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_image(filepath):
    image = Image.open(filepath)
    # Convert image to numpy array
    image_np = np.array(image)
    # Extract text from image
    extracted_text = reader.readtext(image_np)

    # Check if extracted_text is empty
    if not extracted_text:
        return ""

    # Iterate over each element in extracted_text and extract the text content
    text_list = []
    for text_info in extracted_text:
        text_list.append(text_info[1])

    # Join the extracted text content into a single string
    text = '\n'.join(text_list)
    return text

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def store_file_in_database(filename, content):
    cursor = conn.cursor()
    query = "INSERT INTO uploaded_files (filename, content) VALUES (%s, %s)"
    cursor.execute(query, (filename, content))
    conn.commit()
    cursor.close()

def combine_text_files(folder):
    combined_text = ""
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('.txt'):
                file_path = os.path.join(root, file)
                with open(file_path, 'r', encoding='utf-8') as f:
                    combined_text += f.read() + "\n\n"
    return combined_text

if __name__ == '__main__':
    app.run(debug=True)








